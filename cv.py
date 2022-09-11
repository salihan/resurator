from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_LINE_SPACING

#-------------------
from docx.oxml import parse_xml, register_element_cls
from docx.oxml.ns import nsdecls
from docx.oxml.shape import CT_Picture
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne

# refer to docx.oxml.shape.CT_Inline
class CT_Anchor(BaseOxmlElement):
    """
    ``<w:anchor>`` element, container for a floating image.
    """
    extent = OneAndOnlyOne('wp:extent')
    docPr = OneAndOnlyOne('wp:docPr')
    graphic = OneAndOnlyOne('a:graphic')

    @classmethod
    def new(cls, cx, cy, shape_id, pic, pos_x, pos_y):
        """
        Return a new ``<wp:anchor>`` element populated with the values passed
        as parameters.
        """
        anchor = parse_xml(cls._anchor_xml(pos_x, pos_y))
        anchor.extent.cx = cx
        anchor.extent.cy = cy
        anchor.docPr.id = shape_id
        anchor.docPr.name = 'Picture %d' % shape_id
        anchor.graphic.graphicData.uri = (
            'http://schemas.openxmlformats.org/drawingml/2006/picture'
        )
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def new_pic_anchor(cls, shape_id, rId, filename, cx, cy, pos_x, pos_y):
        """
        Return a new `wp:anchor` element containing the `pic:pic` element
        specified by the argument values.
        """
        pic_id = 0  # Word doesn't seem to use this, but does not omit it
        pic = CT_Picture.new(pic_id, filename, rId, cx, cy)
        anchor = cls.new(cx, cy, shape_id, pic, pos_x, pos_y)
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def _anchor_xml(cls, pos_x, pos_y):
        return (
            '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" \n'
            '           behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" \n'
            '           %s>\n'
            '  <wp:simplePos x="0" y="0"/>\n'
            '  <wp:positionH relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionH>\n'
            '  <wp:positionV relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionV>\n'                    
            '  <wp:extent cx="914400" cy="914400"/>\n'
            '  <wp:wrapNone/>\n'
            '  <wp:docPr id="666" name="unnamed"/>\n'
            '  <wp:cNvGraphicFramePr>\n'
            '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
            '  </wp:cNvGraphicFramePr>\n'
            '  <a:graphic>\n'
            '    <a:graphicData uri="URI not set"/>\n'
            '  </a:graphic>\n'
            '</wp:anchor>' % ( nsdecls('wp', 'a', 'pic', 'r'), int(pos_x), int(pos_y) )
        )


# refer to docx.parts.story.BaseStoryPart.new_pic_inline
def new_pic_anchor(part, image_descriptor, width, height, pos_x, pos_y):
    """Return a newly-created `w:anchor` element.

    The element contains the image specified by *image_descriptor* and is scaled
    based on the values of *width* and *height*.
    """
    rId, image = part.get_or_add_image(image_descriptor)
    cx, cy = image.scaled_dimensions(width, height)
    shape_id, filename = part.next_id, image.filename
    return CT_Anchor.new_pic_anchor(shape_id, rId, filename, cx, cy, pos_x, pos_y)


# refer to docx.text.run.add_picture
def add_float_picture(p, image_path_or_stream, width=None, height=None, pos_x=0, pos_y=0):
    """Add float picture at fixed position `pos_x` and `pos_y` to the top-left point of page.
    """
    run = p.add_run()
    anchor = new_pic_anchor(run.part, image_path_or_stream, width, height, pos_x, pos_y)
    run._r.add_drawing(anchor)

# refer to docx.oxml.shape.__init__.py
register_element_cls('wp:anchor', CT_Anchor)
#-------------------

def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

def make_rows_italic(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.italic = True
                    run.font.size = Pt(8)


if __name__ == '__main__':

    name = "Mohd Salihan"
    email = "salihan@yahoo.com"
    phone = "054 512 7917"

    document = Document()

    # profile
    ## add a floating image
    p = document.add_paragraph()
    add_float_picture(p, 'me30percent.jpg', width=Inches(1.25), pos_x=Pt(40), pos_y=Pt(20))

    # document.add_picture('me30percent.jpg', width=Inches(1.25))
    document.add_paragraph(name + ' | ' + email + ' | ' + phone, style='Intense Quote')

    # some heading and para
    document.add_heading('Profile', level=1)
    p = document.add_paragraph('Experienced Software Developer who is open to new challenges, '
                               'with a can-do mind-set and a strong drive to deal with any situations '
                               'that arise in an IT environment. Fast learner who can keep up with any tool '
                               'or methodology such as SDLC or SCRUM agile approach. Able to manage conferences, '
                               'lead a team, do research work, and deliver tasks in a timely manner. '
                               'Able to work effectively either in a team or independently and '
                               'maintain professionalism under pressure.')
    p.alignment = 3  # for left, 1 for center, 2 right, 3 justify
    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True

    # Education
    document.add_heading('Education', level=1)
    document.add_heading('Universiti Teknologi Petronas', level=2)
    document.add_paragraph('Individual final year project: '
                           'Simplifying Object-Oriented Java Programming Language', style='List Bullet')
    document.add_paragraph('Contribution: Shorter reusable codes and a faster system performance compared'
                           ' to the regular native codes.', style='List Bullet')

    # Experiences
    document.add_heading('Experiences', level=1)
    document.add_heading('Software Developer | Jul 2019 – Apr 2021', level=2)
    document.add_heading('BIT Group – Software Dept, Cyberjaya, Malaysia', level=3)
    document.add_paragraph('Developed business intelligence module for Marine Department of Malaysia:')
    document.add_paragraph('Liaised with customers and handled customer requests.', style='List Bullet')
    document.add_paragraph('Assisted in data integration from multiple sources to a single reliable Marine Department of Malaysia MySQL database.', style='List Bullet')
    document.add_paragraph('Handled and resolved data cleaning and preparation by using ETL tool Pentaho Kettle.', style='List Bullet')
    document.add_paragraph('Updated analytical process OLAP and data warehouse deployment.', style='List Bullet')
    document.add_paragraph('Created dashboards and reports for Marine Department of Malaysia.', style='List Bullet')

    # References
    document.add_heading('References')
    # ref1 = ['Higher College of Technology Dubai\n', 'Phone: +9712 2064772\n', 'Email: aamin@hct.ac.ae\n']
    # ref2 = ['Universiti Kebangsaan Malaysia\n', 'Phone: +6012 392 3755\n', 'Email: nazlina.ali@ukm.edu.my\n']
    table = document.add_table(rows=1, cols=2)

    # commented above feature because it's not working. Will solve it later insyaallah
    # tbl_heading = table.rows[0].cells
    # tbl_heading[0].text = 'Dr. Anang Hudaya Muhamad Amin'
    # tbl_heading[1].text = 'Dr. Nazlena Mohamad Ali'
    row_cells = table.add_row().cells
    # row_cells[0].text = ' '.join(ref1[0])
    # row_cells[1].text = ' '.join(ref2[0])
    # row_cells[1].text = ' satu\ndua\n'

    row_cells[0].paragraphs[0].add_run('Dr. Anang Hudaya Muhamad Amin').bold = True
    row_cells[0].paragraphs[0].add_run('\nHigher College of Technology Dubai').italic = True
    row_cells[0].paragraphs[0].add_run('\nPhone: +9712 2064772')
    row_cells[0].paragraphs[0].add_run('\nEmail: aamin@hct.ac.ae')

    row_cells[1].paragraphs[0].add_run('Dr. Nazlena Mohamad Ali').bold = True
    row_cells[1].paragraphs[0].add_run('\nHigher College of Technology Dubai').italic = True
    row_cells[1].paragraphs[0].add_run('\nPhone: +6012 392 3755')
    row_cells[1].paragraphs[0].add_run('\nEmail: nazlina.ali@ukm.edu.my')
    # row_cells = table.add_row().cells
    # row_cells[0].text = ' '.join(ref1[1:])
    # row_cells[1].text = ' '.join(ref2[1:])
    # make_rows_bold(table.rows[0])
    # make_rows_italic(table.rows[1])


    # some records in a table
    # records = (
    #     (3, '101', 'Spam'),
    #     (7, '422', 'Eggs'),
    #     (4, '631', 'Spam, spam, eggs, and spam')
    # )
    # table = document.add_table(rows=1, cols=3)
    # hd_cells = table.rows[0].cells
    # hd_cells[0].text = ' Qty '
    # hd_cells[1].text = ' id '
    # hd_cells[2].text = ' Desc '
    # for qty, id, desc in records:
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = str(qty)
    #     row_cells[1].text = id
    #     row_cells[2].text = desc


    document.save('cv.docx')